# -*- coding=utf-8 -*-
# author = "tungtt"

import os
from docx import Document
from glob import glob
from collections import defaultdict
import os
from io import open


def check_duplicate_answer(data):
    dup_ans = []
    for i,qa in enumerate(data):
        ans = qa[2]
        ans_list = ans.split('\n')
        if len(ans_list)>4:
            continue
        if u'' in ans_list or u' ' in ans_list:
            continue
        if (len(ans_list)) != len(set(ans_list)):
            print("Duplicate answer ")
            print ans
            print ("==========>" + data[i][0])
            ans = ans.replace("\n" , "<br>")
            dup_ans.append(data[i][0] + "<br>" + ans)

    return dup_ans

def list_duplicates(seq):
    tally = defaultdict(list)
    for i, item in enumerate(seq):
        tally[item].append(i)
    return ((key, locs) for key, locs in tally.items()
            if len(locs) > 1)

def check_dup_code(data):
    ma_mon = []
    dup_code = []
    for mm in data:
        ma_mon.append(mm[0])
    for dup in sorted(list_duplicates(ma_mon)):
        indexs = dup[1]
        dup_code.append(data[indexs[0]][0])
        dup_code.append(data[indexs[0]][0])

    print(dup_code)

    return dup_code





def check_duplicate_ques(data):
    dup_quess = []
    question = []
    for dt in data:
        question.append(dt[1].replace('\n',""))
    #print question
    for dup in sorted(list_duplicates(question)):
        dup_q = ""
        indexs = dup[1]
        for idx in indexs:
            print("\n Duplicate questions: ")
            print(data[idx][0])
            print(data[idx][1])
            dup_q += data[idx][0] + "<br>" + data[idx][1] + "<br><br>"
        dup_quess.append(dup_q)
    return dup_quess


        # #comment cho nay
        #     ans.append(data[idx][2].replace('\n',""))
        #
        # for dup_ans in sorted(list_duplicates(ans)):
        #     idxs2 = dup_ans[1]
        #     print("\n Duplicate questions and answers : ")
        #     for idx in idxs2:
        #         ans_idx = indexs[idx]
        #         print data[ans_idx][0]
        #         print(data[ans_idx][1])
        #         print(data[ans_idx][2])

def checking(doc_file):
    print('_'*100)
    print("Process  "+doc_file+'\n')
    doc = Document(doc_file)
    print("Number of table %d" %len(doc.tables))
    table = doc.tables[2]
    data = []

    m_idx = 1
    q_idx = 2
    a_idx = 4

    xx = table.rows

    for i,cont in enumerate(table.rows[0].cells):
        str = cont.text
        if cont.text == u"Nội dung câu hỏi":
            q_idx = i
            break
    for i, cont in enumerate(table.rows[0].cells):
        if cont.text == u"Mã câu":
            m_idx = i
            break

    for i, cont in enumerate(table.rows[0].cells):
        if cont.text == u"Các phương án trả lời":
            a_idx = i
            break

    for row in table.rows:
        x = row.cells
        try:
            ma_mon = x[m_idx].text
        except:
            print("m_idx %d "%m_idx)
        question = x[q_idx].text.lower()
        answer = x[a_idx].text.lower()
        data.append([ma_mon,question,answer])
    dup_code = check_dup_code(data)
    dup_ques = check_duplicate_ques(data)
    dup_ans = check_duplicate_answer(data)
    return dup_code,dup_ques,dup_ans


if __name__ == "__main__":
    checking()
